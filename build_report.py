"""Build the ready-to-submit Assignment 2 report from the verified demo results."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "report" / "Assignment2_Report.docx"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "595959"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALL_OUT = "F4F6F9"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def set_run_font(run, size: float, color: str = "000000", bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_table_geometry(table, widths: list[int], header_fill: str | None = LIGHT_GRAY) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table_pr = table._tbl.tblPr
    layout = table_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    table_width = table_pr.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_pr.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")
    indent = table_pr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        table_pr.append(indent)
    indent.set(qn("w:w"), str(TABLE_INDENT))
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            cell.width = Inches(widths[col_index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_width = tc_pr.first_child_found_in("w:tcW")
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                tc_pr.append(tc_width)
            tc_width.set(qn("w:w"), str(widths[col_index]))
            tc_width.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_index == 0 and header_fill:
                shade(cell, header_fill)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Information Retrieval Assignment 2  |  Research Discovery Hub")
    set_run_font(run, 8.5, MUTED)


def add_title_block(doc: Document) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(16)
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("INFORMATION RETRIEVAL — ASSIGNMENT 2")
    set_run_font(run, 23, "000000", bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run("Research Discovery Hub: an end-to-end Streamlit information retrieval system")
    set_run_font(run, 14, MUTED)
    details = [
        ("Course", "Information Retrieval (AIMLCZG537 / DSECLZG537)"),
        ("Submission", "Assignment 2 — 2025–26, Semester 2"),
        ("Student", "[Replace with your name and BITS ID]"),
        ("Execution", "Streamlit front end; tested with the supplied reproducible corpus"),
    ]
    for label, value in details:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label}: ")
        set_run_font(label_run, 11, "000000", bold=True)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, 11, "000000")
    doc.add_paragraph()


def add_paragraph(doc: Document, text: str, lead: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    if lead:
        run = paragraph.add_run(lead)
        set_run_font(run, 11, "000000", bold=True)
    run = paragraph.add_run(text)
    set_run_font(run, 11)


def add_note(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH], header_fill=None)
    cell = table.cell(0, 0)
    shade(cell, CALL_OUT)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(f"{title}: ")
    set_run_font(run, 10.5, INK, bold=True)
    run = paragraph.add_run(text)
    set_run_font(run, 10.5, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        paragraph = table.cell(0, index).paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(header)
        set_run_font(run, 9.5, INK, bold=True)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(run, 9.5)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build() -> None:
    OUT.parent.mkdir(exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_title_block(doc)

    doc.add_heading("1. Objective and selected use case", level=1)
    add_paragraph(doc, "The system is a research and technical-article discovery application. It lets a user acquire a heterogeneous collection, prepare text, build and inspect an index, issue advanced queries, view PageRank-aware results, receive Top-K recommendations, and evaluate retrieval quality. The scope is deliberately small enough for a complete BITS Virtual Lab demonstration while retaining the full information-retrieval lifecycle.")
    add_note(doc, "System boundary", "All collection, indexing, searching, ranking, recommendation, evaluation and analytics actions are triggered from the Streamlit interface. SQLite and JSON files are implementation storage, not a separate user-facing backend workflow.")

    doc.add_heading("2. End-to-end design", level=1)
    add_paragraph(doc, "The workflow is: Acquire (CSV, Crossref API, or web crawl) → normalize and deduplicate → store metadata and content separately → preprocess and profile text → build inverted index and link graph → BM25 retrieval and PageRank blending → content/hybrid recommendation → qrels-based evaluation and performance analytics.")
    add_paragraph(doc, "Collection is stored in SQLite. Each document has URL, title, source, author, date, category, keywords, raw text and clean text fields. A separate links table stores crawl edges. The index snapshot stores term postings, document lengths and PageRank scores, making index state downloadable and inspectable.")

    doc.add_heading("3. Assignment requirement coverage", level=1)
    add_table(doc,
              ["Requirement", "Implementation evidence"],
              [
                  ["Streamlit workflow", "Dashboard, acquisition, index management, search/ranking, recommendation, evaluation, mining and analytics pages."],
                  ["Heterogeneous acquisition", "CSV upload, Crossref public API and robots-aware multi-seed HTML crawling."],
                  ["Crawling + duplicates", "Configurable depth/page limit; normalized URLs, SHA-256 content hashes and trigram-shingle Jaccard near-duplicate detection."],
                  ["Text mining", "Tokenization, stop-word removal, stemming comparison, keyword extraction, document category/profile and feature-distribution charts."],
                  ["Search and ranking", "Phrase/AND/OR/NOT filtering, inverted-index candidates, BM25, iterative PageRank and an adjustable blend explanation chart."],
                  ["Recommendation", "TF-IDF cosine Top-K content scores plus an optional feedback-aware hybrid component."],
                  ["Evaluation", "Precision, Recall, F1, P@K, R@K, MAP, MRR and NDCG@K, reported for BM25 and BM25 + PageRank."],
              ], [2700, 6660])

    doc.add_heading("4. Corpus and experimental method", level=1)
    add_paragraph(doc, "The supplied demonstration collection contains 15 compact technical documents and a reproducible citation graph. The first verified index build contained 395 stemmed vocabulary terms with an average document length of 48.3 terms. Five queries have graded qrels: information retrieval, BM25 ranking, duplicate, recommendation similarity and ranking. Evaluation uses K = 5.")
    add_paragraph(doc, "For the ranking experiment, the application first filters candidates using the inverted index and Boolean/phrase constraints. It then applies BM25. A normalized PageRank score from the stored link graph is blended at 20% with the normalized BM25 score. The slider permits a fair comparison with BM25-only ranking at 0% PageRank influence.")
    add_note(doc, "Reproducibility", "Open Acquire & Crawl → Demo dataset → Load or refresh demo corpus, then run Evaluation with the supplied qrels at K = 5. The values below are from that verified run; results will change when the corpus, qrels or weighting is changed.")

    doc.add_heading("5. Experimental results and discussion", level=1)
    add_table(doc,
              ["Strategy", "P@5", "R@5", "F1", "MAP", "MRR", "NDCG@5"],
              [
                  ["BM25", "0.853", "0.530", "0.614", "0.508", "1.000", "0.739"],
                  ["BM25 + PageRank", "0.893", "0.580", "0.659", "0.548", "1.000", "0.790"],
              ], [2520, 1140, 1140, 1140, 1140, 1140, 1140])
    add_paragraph(doc, "The PageRank blend improves the compact demonstration set on P@5 (+0.040), R@5 (+0.050), MAP (+0.040) and NDCG@5 (+0.051). It does not alter MRR because the first relevant result was already first in the evaluated queries. This is a useful result: authority signals help ordering and coverage, but should complement rather than replace lexical relevance.")
    add_table(doc,
              ["Query: ranking", "P@5", "R@5", "MRR", "NDCG@5"],
              [
                  ["BM25", "0.600", "0.750", "1.000", "0.668"],
                  ["BM25 + PageRank", "0.800", "1.000", "1.000", "0.921"],
              ], [3600, 1440, 1440, 1440, 1440])
    add_paragraph(doc, "For the focused ranking query, the PageRank-aware blend moves authoritative ranking material higher, raising NDCG@5 substantially. The application exposes BM25, PageRank and final score next to every result, so the observed change is interpretable rather than a black-box score.")

    doc.add_heading("6. Compulsory inferences and discussion", level=1)
    doc.add_heading("6.1 Relevant documents retrieved but poorly ranked", level=2)
    add_paragraph(doc, "If relevant documents are retrieved but appear low in the list, candidate generation is probably adequate and the ranking signal is weak or misweighted. Likely causes are untuned BM25 k1/b parameters, inaccurate length normalization, no title/metadata field boost, vocabulary mismatch, missing phrase/proximity features, stale link authority, or qrels that do not reflect user intent. Improvements are to tune BM25 on held-out qrels, add field and freshness boosts, use query expansion or synonym handling, blend PageRank conservatively, and train a learning-to-rank re-ranker on judged query-document pairs. Per-query P@K and NDCG diagnostics are more actionable than only one global score.")
    doc.add_heading("6.2 Effect of duplicate and near-duplicate documents", level=2)
    add_paragraph(doc, "Duplicates bloat postings and storage, may occupy several high-ranked positions, and reduce result diversity. They bias content recommendations toward repeated copies and can inflate evaluation when many copies are judged relevant. The system mitigates exact URL and content duplicates with URL normalization and a SHA-256 normalized-content hash. It flags near duplicates with trigram shingles and Jaccard similarity. At larger scale, MinHash/LSH or SimHash should be used for faster approximate duplicate detection, with canonical URLs and cluster-level evaluation judgments.")
    doc.add_heading("6.3 Content-based versus collaborative recommendation", level=2)
    add_paragraph(doc, "Content-based recommendation uses the document itself, here TF-IDF cosine similarity. It is appropriate when item text is rich, explanations matter, or the user and corpus are new; it suffers when a profile becomes too narrow. Collaborative recommendation uses interaction patterns such as saves, ratings and clicks. It can discover useful cross-topic material without rich text, but requires enough users and interactions and has cold-start weaknesses. A hybrid approach is preferable once feedback exists because it preserves content coverage while adding behavioral personalization.")
    doc.add_heading("6.4 Value of the integrated lifecycle", level=2)
    add_paragraph(doc, "Each lifecycle stage protects the next. Responsible crawling and API/CSV acquisition improve coverage; deduplication improves collection quality; preprocessing and indexing make search efficient; ranking orders candidates by relevance and authority; recommendation extends discovery beyond an explicit query; and evaluation turns user-visible behavior into measurable evidence. Keeping these stages in one Streamlit workflow makes failures traceable—for example, a poor result can be investigated through its source, preprocessing, term vocabulary, BM25 score, PageRank value and relevance judgment.")
    doc.add_heading("6.5 Learnings from the experiment", level=2)
    add_paragraph(doc, "The work shows that a good IR system is not a single algorithm. The baseline BM25 index provides a strong transparent starting point, but authority features changed the rank quality for the diagnostic query. The preprocessing comparison shows why feature choices control both vocabulary size and matching behavior. The duplicate checks and separate metadata/content storage improve reliability before a ranking model runs. Finally, the metric dashboard prevents overclaiming: the PageRank blend improved NDCG and recall in this small curated collection, but it must be retested on a larger, independently judged corpus before treating it as a general conclusion.")

    doc.add_heading("7. Virtual Lab demonstration evidence", level=1)
    add_note(doc, "Student action required", "Run the app on the BITS Virtual Lab and replace the three placeholders below with your own screenshots. They cannot be truthfully generated outside your Virtual Lab session.")
    for caption in (
        "Screenshot 1 — Dashboard after loading the demo or chosen dataset. Show corpus count, vocabulary and acquisition-source chart.",
        "Screenshot 2 — Search & Ranking for `ranking`. Show the BM25, PageRank and final scores with PageRank influence at 0.20.",
        "Screenshot 3 — Evaluation Dashboard at K = 5. Show the BM25 versus BM25 + PageRank metric table and chart.",
    ):
        table = doc.add_table(rows=1, cols=1)
        set_table_geometry(table, [TABLE_WIDTH], header_fill=None)
        cell = table.cell(0, 0)
        shade(cell, "FAFAFA")
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(18)
        paragraph.paragraph_format.space_after = Pt(18)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(caption + "\n[Insert Virtual Lab screenshot here]")
        set_run_font(run, 10.5, MUTED, italic=True)
        doc.add_paragraph().paragraph_format.space_after = Pt(3)

    doc.add_heading("8. Execution instructions", level=1)
    add_paragraph(doc, "Create a virtual environment, install requirements, and run `streamlit run app.py`. The README lists every step and the required CSV schemas. For the demonstration, use the supplied data/sample_documents.csv and data/sample_qrels.csv or add a permitted domain/API/CSV collection through the acquisition page. No command other than starting Streamlit is needed for the normal front-end workflow.")
    add_note(doc, "Submission contents", "Submit app.py, requirements.txt, data/sample_documents.csv, data/sample_qrels.csv, README.md, this report after filling student details and screenshots, plus any additional dataset material used in your own run.")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
